import spacy
import language_tool_python
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.core.files.storage import default_storage
from .models import Document, Content
from rest_framework.permissions import IsAuthenticated
from docx import Document as DocxDocument  # For .docx files
import fitz  # PyMuPDF for .pdf files
import os
from django.http import HttpResponse
from django.core.files.storage import default_storage
from docx import Document as DocxDocument
# .......................................................
from .serializers import DocumentSerializer, ContentSerializer
from rest_framework.permissions import IsAuthenticated
import logging
import language_tool_python

from rest_framework.authtoken.views import ObtainAuthToken
from rest_framework.permissions import AllowAny
from rest_framework.authtoken.models import Token

from django.contrib.auth.models import User








logger = logging.getLogger(__name__)
nlp = spacy.load("en_core_web_sm")
tool = language_tool_python.LanguageTool('en-US')

def read_docx(file_path):
    """Extract text from a .docx file."""
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def read_pdf(file_path):
    """Extract text from a .pdf file."""
    pdf_text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            pdf_text += page.get_text()
    return pdf_text


class UploadDocumentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        file = request.FILES.get('document')
        if not file:
            return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)

        # Determine file type based on extension
        file_extension = file.name.split('.')[-1].lower()
        try:
            # Save the file temporarily to read it
            file_path = default_storage.save(file.name, file)
            full_path = default_storage.path(file_path)

            # Read content based on file extension
            if file_extension == 'txt':
                with open(full_path, 'r') as f:
                    content = f.read()
            elif file_extension == 'docx':
                content = read_docx(full_path)
            elif file_extension == 'pdf':
                content = read_pdf(full_path)
            else:
                return Response({'error': 'Unsupported file type'}, status=status.HTTP_400_BAD_REQUEST)

            # Create a Document entry and store original content
            document = Document.objects.create(user=request.user, status='uploaded')
            Content.objects.create(document=document, original_content=content)

            return Response({'id': document.id}, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({'error': 'Failed to process document'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




class GetDocumentView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, id):
        try:
            document = Document.objects.get(id=id, user=request.user)
            content = Content.objects.get(document=document)
            return Response({
                'original_content': content.original_content,
                'improved_content': content.improved_content
            })
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Content.DoesNotExist:
            return Response({'error': 'Content not found'}, status=status.HTTP_404_NOT_FOUND)






import spacy
import language_tool_python
from .nlp_utils import nlp, tool

class ImproveDocumentView(APIView):
    permission_classes = [IsAuthenticated]

    def improve_text(self, text):
        # Tokenize and analyze text with spaCy
        doc = nlp(text)
        
        # Initialize LanguageTool for grammar checking
        matches = tool.check(text)
        improved_text = text

        suggestions = []
        offset = 0  # Track the offset due to replacements

        # Apply grammar suggestions from LanguageTool
        for match in matches:
            suggestion = {
                "text": match.context,
                "error": match.message,
                "suggestions": match.replacements
            }
            suggestions.append(suggestion)

            # Apply the first suggestion if available
            if match.replacements:
                start = match.offset + offset
                end = start + match.errorLength
                replacement = match.replacements[0]
                
                # Replace the text and update the offset
                improved_text = improved_text[:start] + replacement + improved_text[end:]
                offset += len(replacement) - match.errorLength

        # Use spaCy to further improve clarity, style, and structure
        for sentence in doc.sents:
            tokens = [token.text for token in sentence]
            entities = [(ent.text, ent.label_) for ent in sentence.ents]

            # Example: Rephrase sentences with passive voice or improve clarity
            if any(token.dep_ == "auxpass" for token in sentence):
                passive_suggestion = f"Consider rephrasing '{sentence.text}' to active voice."
                suggestions.append({
                    "text": sentence.text,
                    "error": "Passive voice detected",
                    "suggestions": [passive_suggestion]
                })

            # Example: Detect and suggest correction for overly complex sentences
            if len(tokens) > 20:
                clarity_suggestion = f"Consider simplifying the sentence: '{sentence.text}'"
                suggestions.append({
                    "text": sentence.text,
                    "error": "Complex sentence",
                    "suggestions": [clarity_suggestion]
                })

            # Example: Check for redundant expressions (e.g., "advance planning")
            redundant_phrases = ["advance planning", "free gift", "future plans"]
            for phrase in redundant_phrases:
                if phrase in sentence.text.lower():
                    redundancy_suggestion = f"Remove redundant phrase '{phrase}'"
                    suggestions.append({
                        "text": sentence.text,
                        "error": "Redundant expression",
                        "suggestions": [redundancy_suggestion]
                    })

        return improved_text, suggestions

    def post(self, request, id):
        try:
            document = Document.objects.get(id=id, user=request.user)
            content = Content.objects.get(document=document)

            # Generate improved text and suggestions
            improved_text, suggestions = self.improve_text(content.original_content)

            # Save improved text to improved_content
            content.improved_content = improved_text
            content.save()

            return Response({"suggestions": suggestions}, status=status.HTTP_200_OK)

        except Document.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)








class UserRegistrationView(APIView):
    permission_classes = [AllowAny]  # This makes the endpoint accessible to all

    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        if username and password:
            user = User.objects.create_user(username=username, password=password)
            return Response({'id': user.id}, status=status.HTTP_201_CREATED)
        return Response({'error': 'Username and password required'}, status=status.HTTP_400_BAD_REQUEST)



class UserLoginView(ObtainAuthToken):
    def post(self, request):
        serializer = self.serializer_class(data=request.data, context={'request': request})
        serializer.is_valid(raise_exception=True)
        user = serializer.validated_data['user']
        token, created = Token.objects.get_or_create(user=user)
        return Response({'token': token.key})




class ExportDocumentView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, id, format_type):
        try:
            document = Document.objects.get(id=id, user=request.user)
            content = Content.objects.get(document=document)

            # Debugging: Print content to verify it
            print("Improved Content:", content.improved_content)

            if format_type == 'txt':
                response = HttpResponse(content.improved_content, content_type='text/plain')
                response['Content-Disposition'] = f'attachment; filename="improved_document_{id}.txt"'
                return response

            elif format_type == 'docx':
                # Create a .docx file dynamically
                doc = DocxDocument()
                doc.add_paragraph(content.improved_content)
                file_path = f'improved_document_{id}.docx'
                doc.save(file_path)

                # Serve the .docx file as a download
                with open(file_path, 'rb') as f:
                    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename="improved_document_{id}.docx"'
                    os.remove(file_path)  # Clean up the file after download
                    return response

            else:
                return Response({'error': 'Unsupported format'}, status=status.HTTP_400_BAD_REQUEST)

        except Document.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
